#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Генератор отчетов для UX анализатора
"""

from typing import Dict, Any, List
import json
import os
from pathlib import Path

class ReportGenerator:
    """Генератор отчетов в различных форматах"""
    
    def __init__(self, config):
        self.config = config
        self.colors = {
            'primary': '#2563eb',
            'secondary': '#7c3aed',
            'success': '#10b981',
            'warning': '#f59e0b',
            'danger': '#ef4444',
            'dark': '#1f2937',
            'light': '#f3f4f6',
            'white': '#ffffff',
            'text': '#111827',
            'text_secondary': '#6b7280',
            'border': '#e5e7eb',
            'background': '#f9fafb'
        }
    
    def generate_html(self, analysis_result) -> str:
        """Генерация HTML отчета"""
        
        # Подготовка данных
        total_interviews = analysis_result.total_interviews
        duration = analysis_result.analysis_duration
        api_calls = analysis_result.api_calls
        total_cost = analysis_result.total_cost
        
        # Извлечение ключевых данных
        key_insights = []
        pain_points = []
        recommendations = []
        behavioral_patterns = []
        
        if analysis_result.research_findings:
            findings = analysis_result.research_findings
            key_insights = findings.key_insights or []
            recommendations = findings.recommendations or []
            behavioral_patterns = findings.behavioral_patterns or []
        
        # Сбор всех болей
        for summary in analysis_result.interview_summaries:
            pain_points.extend(summary.pain_points or [])
        
        # Статистика
        total_pains = len(pain_points)
        total_needs = sum(len(s.needs or []) for s in analysis_result.interview_summaries)
        total_quotes = sum(len(s.quotes or []) for s in analysis_result.interview_summaries)
        
        html = f"""<!DOCTYPE html>
<html lang="ru">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>UX Research Report - {self.config.report_title}</title>
    <link href="https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800;900&display=swap" rel="stylesheet">
    <style>
        {self._get_css_styles()}
    </style>
</head>
<body>
    <div class="container">
        <!-- Заголовок -->
        {self._generate_cover_page()}
        
        <!-- Основные метрики -->
        <section class="metrics-section">
            <h2>📈 Ключевые метрики</h2>
            <div class="metrics-grid">
                <div class="metric-card">
                    <div class="metric-value">{total_interviews}</div>
                    <div class="metric-label">Интервью проанализировано</div>
                </div>
                <div class="metric-card">
                    <div class="metric-value">{total_pains}</div>
                    <div class="metric-label">Точек боли выявлено</div>
                </div>
                <div class="metric-card">
                    <div class="metric-value">{total_needs}</div>
                    <div class="metric-label">Потребностей выявлено</div>
                </div>
                <div class="metric-card">
                    <div class="metric-value">{total_quotes}</div>
                    <div class="metric-label">Цитат извлечено</div>
                </div>
            </div>
        </section>

        <!-- Ключевые инсайты -->
        {self._generate_insights_section(key_insights)}

        <!-- Детальные инсайты -->
        {self._generate_insights_section_full(key_insights)}

        <!-- Точки боли -->
        {self._generate_pain_points_section(pain_points)}

        <!-- Детальный анализ точек боли -->
        {self._generate_pain_points_section_full(pain_points)}

        <!-- Поведенческие паттерны -->
        {self._generate_behavioral_patterns_section_full(behavioral_patterns)}

        <!-- Эмоциональное путешествие -->
        {self._generate_emotional_journey_section(analysis_result.interview_summaries)}

        <!-- Противоречия -->
        {self._generate_contradictions_section(analysis_result.interview_summaries)}

        <!-- Цитаты -->
        {self._generate_quotes_section(analysis_result.interview_summaries)}

        <!-- Рекомендации -->
        {self._generate_recommendations_section(recommendations)}

        <!-- Детальные рекомендации -->
        {self._generate_recommendations_section_full(recommendations)}

        <!-- Матрица приоритетов -->
        {self._generate_priority_matrix_section(recommendations)}

        <!-- Дорожная карта -->
        {self._generate_roadmap_section({})}

        <!-- Детали по интервью -->
        {self._generate_interviews_section(analysis_result.interview_summaries)}

        <!-- Футер -->
        <footer class="footer">
            <p>Отчет сгенерирован автоматически с помощью UX Analyzer V25.0</p>
            <p>Модель: Claude 3.5 Sonnet | API: OpenRouter</p>
        </footer>
    </div>
</body>
</html>"""
        
        return html
    
    def generate_docx(self, analysis_result) -> str:
        """Генерация DOCX отчета"""
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            doc = Document()
            
            # Заголовок
            title = doc.add_heading('UX Research Report', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Мета-информация
            doc.add_paragraph(f'Интервью проанализировано: {analysis_result.total_interviews}')
            doc.add_paragraph(f'Время анализа: {analysis_result.analysis_duration:.1f} сек')
            doc.add_paragraph(f'API вызовов: {analysis_result.api_calls}')
            doc.add_paragraph(f'Примерная стоимость: ${analysis_result.total_cost:.4f}')
            
            # Ключевые инсайты
            if analysis_result.research_findings:
                doc.add_heading('Ключевые инсайты', level=1)
                for insight in analysis_result.research_findings.key_insights[:5]:
                    doc.add_heading(insight.get('problem_title', 'Инсайт'), level=2)
                    doc.add_paragraph(insight.get('problem_description', 'Описание недоступно'))
            
            # Сохранение
            output_file = "ux_analysis_report.docx"
            doc.save(output_file)
            return output_file
            
        except ImportError:
            print("❌ Для генерации DOCX требуется python-docx")
            return ""
        except Exception as e:
            print(f"❌ Ошибка при генерации DOCX: {e}")
            return ""
    
    def generate_pdf(self, html_content: str) -> str:
        """Генерация PDF отчета"""
        try:
            from weasyprint import HTML
            
            output_file = "ux_analysis_report.pdf"
            HTML(string=html_content).write_pdf(output_file)
            return output_file
            
        except ImportError:
            print("❌ Для генерации PDF требуется weasyprint")
            return ""
        except Exception as e:
            print(f"❌ Ошибка при генерации PDF: {e}")
            return ""
    
    def _generate_cover_page(self) -> str:
        """Генерация титульной страницы"""
        return f"""
        <header class="header">
            <div class="header-content">
                <h1>🔬 UX Research Report</h1>
                <p class="subtitle">Анализ интервью с использованием AI</p>
                <div class="company-info">
                    <p class="company-name">{self.config.company_name}</p>
                    <p class="report-date">{self._get_current_date()}</p>
                </div>
            </div>
        </header>
        """
    
    def _get_current_date(self) -> str:
        """Получение текущей даты"""
        from datetime import datetime
        return datetime.now().strftime("%d.%m.%Y")
    
    def _generate_insights_section(self, insights: List[Dict[str, Any]]) -> str:
        """Генерация секции с инсайтами"""
        if not insights:
            return ""
        
        insights_html = '<section class="insights-section">\n'
        insights_html += '<h2>💡 Ключевые инсайты</h2>\n'
        insights_html += '<div class="insights-grid">\n'
        
        for i, insight in enumerate(insights[:6], 1):  # Показываем максимум 6
            title = insight.get('problem_title', f'Инсайт {i}')
            description = insight.get('problem_description', 'Описание недоступно')
            severity = insight.get('severity', 'medium')
            affected = insight.get('affected_percentage', 'N/A')
            
            severity_class = f"severity-{severity}"
            
            insights_html += f'''
            <div class="insight-card {severity_class}">
                <div class="insight-header">
                    <h3>{title}</h3>
                    <span class="severity-badge {severity_class}">{severity.upper()}</span>
                </div>
                <p class="insight-description">{description}</p>
                <div class="insight-meta">
                    <span class="affected-users">👥 {affected}</span>
                </div>
            </div>'''
        
        insights_html += '</div>\n</section>\n'
        return insights_html
    
    def _generate_insights_section_full(self, insights: List[Dict[str, Any]]) -> str:
        """Полная секция с инсайтами"""
        if not insights:
            return ""
        
        insights_html = '<section class="insights-section-full">\n'
        insights_html += '<h2>💡 Детальные инсайты</h2>\n'
        
        # Группировка по серьезности
        severity_groups = {'critical': [], 'high': [], 'medium': [], 'low': []}
        for insight in insights:
            severity = insight.get('severity', 'medium')
            if severity in severity_groups:
                severity_groups[severity].append(insight)
        
        for severity, group_insights in severity_groups.items():
            if not group_insights:
                continue
            
            severity_name = {
                'critical': 'Критические',
                'high': 'Высокие',
                'medium': 'Средние',
                'low': 'Низкие'
            }.get(severity, severity)
            
            insights_html += f'<div class="severity-group">\n'
            insights_html += f'<h3 class="severity-title {severity}">{severity_name} инсайты</h3>\n'
            insights_html += '<div class="insights-detailed-grid">\n'
            
            for insight in group_insights:
                title = insight.get('problem_title', 'Инсайт')
                description = insight.get('problem_description', 'Описание недоступно')
                affected = insight.get('affected_percentage', 'N/A')
                priority = insight.get('priority', 'N/A')
                effort = insight.get('effort', 'N/A')
                
                insights_html += f'''
                <div class="insight-detailed-card {severity}">
                    <div class="insight-detailed-header">
                        <h4>{title}</h4>
                        <div class="insight-badges">
                            <span class="priority-badge priority-{priority.lower()}">{priority}</span>
                            <span class="effort-badge effort-{effort.lower()}">{effort}</span>
                        </div>
                    </div>
                    <p class="insight-detailed-description">{description}</p>
                    <div class="insight-detailed-meta">
                        <span class="affected-users">👥 {affected}</span>
                    </div>
                </div>'''
            
            insights_html += '</div>\n</div>\n'
        
        insights_html += '</section>\n'
        return insights_html
    
    def _generate_pain_points_section(self, pain_points: List[Dict[str, Any]]) -> str:
        """Генерация секции с точками боли"""
        if not pain_points:
            return ""
        
        # Группировка по серьезности
        severity_groups = {'critical': [], 'high': [], 'medium': [], 'low': []}
        for pain in pain_points:
            severity = pain.get('severity', 'medium')
            if severity in severity_groups:
                severity_groups[severity].append(pain)
        
        pain_html = '<section class="pain-points-section">\n'
        pain_html += '<h2>😤 Точки боли пользователей</h2>\n'
        
        for severity, pains in severity_groups.items():
            if not pains:
                continue
            
            severity_name = {
                'critical': 'Критические',
                'high': 'Высокие',
                'medium': 'Средние',
                'low': 'Низкие'
            }.get(severity, severity)
            
            pain_html += f'<div class="severity-group">\n'
            pain_html += f'<h3 class="severity-title {severity}">{severity_name} проблемы</h3>\n'
            pain_html += '<div class="pain-points-grid">\n'
            
            for pain in pains[:4]:  # Максимум 4 на группу
                pain_text = pain.get('pain', 'Описание недоступно')
                context = pain.get('context', '')
                
                pain_html += f'''
                <div class="pain-point-card {severity}">
                    <h4>{pain_text[:100]}{"..." if len(pain_text) > 100 else ""}</h4>
                    {f'<p class="pain-context">{context[:150]}{"..." if len(context) > 150 else ""}</p>' if context else ''}
                </div>'''
            
            pain_html += '</div>\n</div>\n'
        
        pain_html += '</section>\n'
        return pain_html
    
    def _generate_pain_points_section_full(self, problems: List[Dict[str, Any]], charts: str = "") -> str:
        """Полная секция с точками боли"""
        if not problems:
            return ""
        
        pain_html = '<section class="pain-points-section-full">\n'
        pain_html += '<h2>😤 Детальный анализ точек боли</h2>\n'
        
        if charts:
            pain_html += f'<div class="charts-container">{charts}</div>\n'
        
        # Группировка по типу боли
        type_groups = {}
        for problem in problems:
            pain_type = problem.get('pain_type', 'other')
            if pain_type not in type_groups:
                type_groups[pain_type] = []
            type_groups[pain_type].append(problem)
        
        for pain_type, type_problems in type_groups.items():
            type_name = {
                'functional': 'Функциональные',
                'process': 'Процессные',
                'emotional': 'Эмоциональные',
                'social': 'Социальные',
                'financial': 'Финансовые',
                'other': 'Другие'
            }.get(pain_type, pain_type.title())
            
            pain_html += f'<div class="pain-type-group">\n'
            pain_html += f'<h3 class="pain-type-title">{type_name} проблемы</h3>\n'
            pain_html += '<div class="pain-points-detailed-grid">\n'
            
            for problem in type_problems:
                pain_text = problem.get('pain', 'Описание недоступно')
                context = problem.get('context', '')
                severity = problem.get('severity', 'medium')
                frequency = problem.get('frequency', 'N/A')
                impact = problem.get('impact', '')
                
                pain_html += f'''
                <div class="pain-point-detailed-card {severity}">
                    <div class="pain-point-detailed-header">
                        <h4>{pain_text[:100]}{"..." if len(pain_text) > 100 else ""}</h4>
                        <span class="severity-badge {severity}">{severity.upper()}</span>
                    </div>
                    {f'<p class="pain-context">{context[:150]}{"..." if len(context) > 150 else ""}</p>' if context else ''}
                    <div class="pain-point-meta">
                        <span class="frequency">🔄 {frequency}</span>
                        {f'<span class="impact">💥 {impact}</span>' if impact else ''}
                    </div>
                </div>'''
            
            pain_html += '</div>\n</div>\n'
        
        pain_html += '</section>\n'
        return pain_html
    
    def _generate_behavioral_patterns_section_full(self, patterns: List[Dict[str, Any]]) -> str:
        """Полная секция с поведенческими паттернами"""
        if not patterns:
            return ""
        
        patterns_html = '<section class="behavioral-patterns-section">\n'
        patterns_html += '<h2>🔄 Поведенческие паттерны</h2>\n'
        patterns_html += '<div class="patterns-grid">\n'
        
        for pattern in patterns:
            pattern_desc = pattern.get('pattern', 'Описание недоступно')
            pattern_type = pattern.get('pattern_type', 'N/A')
            frequency = pattern.get('frequency', 'N/A')
            confidence = pattern.get('confidence', 'N/A')
            underlying_need = pattern.get('underlying_need', '')
            design_implication = pattern.get('design_implication', '')
            
            patterns_html += f'''
            <div class="pattern-card">
                <div class="pattern-header">
                    <h3>Паттерн: {pattern_type.title()}</h3>
                    <div class="pattern-badges">
                        <span class="frequency-badge">{frequency}</span>
                        <span class="confidence-badge confidence-{confidence}">{confidence}</span>
                    </div>
                </div>
                <p class="pattern-description">{pattern_desc}</p>
                {f'<p class="underlying-need"><strong>Глубинная потребность:</strong> {underlying_need}</p>' if underlying_need else ''}
                {f'<p class="design-implication"><strong>Дизайн-импликация:</strong> {design_implication}</p>' if design_implication else ''}
            </div>'''
        
        patterns_html += '</div>\n</section>\n'
        return patterns_html
    
    def _generate_emotional_journey_section(self, summaries: List[Any]) -> str:
        """Секция эмоционального путешествия"""
        if not summaries:
            return ""
        
        # Сбор всех эмоциональных моментов
        all_emotions = []
        for summary in summaries:
            all_emotions.extend(summary.emotional_journey or [])
        
        if not all_emotions:
            return ""
        
        emotions_html = '<section class="emotional-journey-section">\n'
        emotions_html += '<h2>😊 Эмоциональное путешествие пользователей</h2>\n'
        emotions_html += '<div class="emotional-journey-timeline">\n'
        
        # Группировка по валентности
        valence_groups = {'positive': [], 'negative': [], 'mixed': []}
        for emotion in all_emotions:
            valence = emotion.get('valence', 'mixed')
            if valence in valence_groups:
                valence_groups[valence].append(emotion)
        
        for valence, emotions in valence_groups.items():
            if not emotions:
                continue
            
            valence_name = {
                'positive': 'Положительные',
                'negative': 'Отрицательные',
                'mixed': 'Смешанные'
            }.get(valence, valence)
            
            emotions_html += f'<div class="valence-group {valence}">\n'
            emotions_html += f'<h3 class="valence-title">{valence_name} эмоции</h3>\n'
            emotions_html += '<div class="emotions-list">\n'
            
            for emotion in emotions[:5]:  # Максимум 5 на группу
                moment = emotion.get('moment', '')
                emotion_name = emotion.get('emotion', '')
                intensity = emotion.get('intensity', 5)
                quote = emotion.get('quote', '')
                
                emotions_html += f'''
                <div class="emotion-item {valence}">
                    <div class="emotion-header">
                        <h4>{emotion_name}</h4>
                        <div class="intensity-bar">
                            <div class="intensity-fill" style="width: {intensity * 10}%"></div>
                        </div>
                    </div>
                    <p class="emotion-moment">{moment}</p>
                    {f'<blockquote class="emotion-quote">{quote[:200]}{"..." if len(quote) > 200 else ""}</blockquote>' if quote else ''}
                </div>'''
            
            emotions_html += '</div>\n</div>\n'
        
        emotions_html += '</div>\n</section>\n'
        return emotions_html
    
    def _generate_contradictions_section(self, summaries: List[Any]) -> str:
        """Секция противоречий"""
        if not summaries:
            return ""
        
        # Сбор всех противоречий
        all_contradictions = []
        for summary in summaries:
            all_contradictions.extend(summary.contradictions or [])
        
        if not all_contradictions:
            return ""
        
        contradictions_html = '<section class="contradictions-section">\n'
        contradictions_html += '<h2>⚠️ Противоречия в данных</h2>\n'
        contradictions_html += '<div class="contradictions-grid">\n'
        
        for contradiction in all_contradictions:
            contradiction_type = contradiction.get('contradiction_type', 'N/A')
            severity = contradiction.get('severity', 'medium')
            nature = contradiction.get('analysis', {}).get('nature', '')
            
            contradictions_html += f'''
            <div class="contradiction-card {severity}">
                <div class="contradiction-header">
                    <h3>Противоречие: {contradiction_type.title()}</h3>
                    <span class="severity-badge {severity}">{severity.upper()}</span>
                </div>
                <p class="contradiction-nature">{nature}</p>
                <div class="contradiction-statements">
                    <div class="statement">
                        <strong>Утверждение 1:</strong> {contradiction.get('statement_1', {}).get('text', 'N/A')}
                    </div>
                    <div class="statement">
                        <strong>Утверждение 2:</strong> {contradiction.get('statement_2', {}).get('text', 'N/A')}
                    </div>
                </div>
            </div>'''
        
        contradictions_html += '</div>\n</section>\n'
        return contradictions_html
    
    def _generate_quotes_section(self, summaries: List[Any]) -> str:
        """Секция важных цитат"""
        if not summaries:
            return ""
        
        # Сбор всех цитат
        all_quotes = []
        for summary in summaries:
            all_quotes.extend(summary.quotes or [])
        
        if not all_quotes:
            return ""
        
        quotes_html = '<section class="quotes-section">\n'
        quotes_html += '<h2>💬 Ключевые цитаты пользователей</h2>\n'
        quotes_html += '<div class="quotes-grid">\n'
        
        for quote in all_quotes[:10]:  # Максимум 10 цитат
            text = quote.get('text', '')
            context = quote.get('context', '')
            significance = quote.get('significance', '')
            quote_type = quote.get('quote_type', 'general')
            
            quotes_html += f'''
            <div class="quote-card {quote_type}">
                <blockquote class="quote-text">"{text}"</blockquote>
                <div class="quote-meta">
                    <span class="quote-type">{quote_type.title()}</span>
                    {f'<p class="quote-context">{context}</p>' if context else ''}
                    {f'<p class="quote-significance">{significance}</p>' if significance else ''}
                </div>
            </div>'''
        
        quotes_html += '</div>\n</section>\n'
        return quotes_html
    
    def _generate_recommendations_section(self, recommendations: List[Dict[str, Any]]) -> str:
        """Генерация секции с рекомендациями"""
        if not recommendations:
            return ""
        
        recs_html = '<section class="recommendations-section">\n'
        recs_html += '<h2>🚀 Рекомендации</h2>\n'
        recs_html += '<div class="recommendations-grid">\n'
        
        for i, rec in enumerate(recommendations[:6], 1):  # Максимум 6
            title = rec.get('title', f'Рекомендация {i}')
            description = rec.get('description', 'Описание недоступно')
            priority = rec.get('priority', 'P2')
            effort = rec.get('effort', 'M')
            
            recs_html += f'''
            <div class="recommendation-card">
                <div class="rec-header">
                    <h3>{title}</h3>
                    <div class="rec-badges">
                        <span class="priority-badge priority-{priority.lower()}">{priority}</span>
                        <span class="effort-badge effort-{effort.lower()}">{effort}</span>
                    </div>
                </div>
                <p class="rec-description">{description}</p>
            </div>'''
        
        recs_html += '</div>\n</section>\n'
        return recs_html
    
    def _generate_recommendations_section_full(self, recommendations: List[Dict[str, Any]]) -> str:
        """Полная секция с рекомендациями"""
        if not recommendations:
            return ""
        
        recs_html = '<section class="recommendations-section-full">\n'
        recs_html += '<h2>🚀 Стратегические рекомендации</h2>\n'
        recs_html += '<div class="recommendations-detailed-grid">\n'
        
        for rec in recommendations:
            title = rec.get('recommendation', 'Рекомендация')
            rationale = rec.get('rationale', '')
            expected_outcome = rec.get('expected_outcome', '')
            timeline = rec.get('timeline', '')
            success_metrics = rec.get('success_metrics', [])
            
            recs_html += f'''
            <div class="recommendation-detailed-card">
                <div class="rec-detailed-header">
                    <h3>{title}</h3>
                </div>
                {f'<p class="rec-rationale"><strong>Обоснование:</strong> {rationale}</p>' if rationale else ''}
                {f'<p class="rec-outcome"><strong>Ожидаемый результат:</strong> {expected_outcome}</p>' if expected_outcome else ''}
                {f'<p class="rec-timeline"><strong>Сроки:</strong> {timeline}</p>' if timeline else ''}
                {f'<div class="rec-metrics"><strong>Метрики успеха:</strong><ul>{chr(10).join([f"<li>{metric}</li>" for metric in success_metrics])}</ul></div>' if success_metrics else ''}
            </div>'''
        
        recs_html += '</div>\n</section>\n'
        return recs_html
    
    def _generate_priority_matrix_section(self, recommendations: List[Dict[str, Any]]) -> str:
        """Секция матрицы приоритетов"""
        if not recommendations:
            return ""
        
        matrix_html = '<section class="priority-matrix-section">\n'
        matrix_html += '<h2>📊 Матрица приоритетов</h2>\n'
        matrix_html += '<div class="priority-matrix">\n'
        
        # Создание матрицы 2x2
        matrix_html += '''
        <div class="matrix-grid">
            <div class="matrix-cell high-impact-high-effort">
                <h4>Высокое влияние, Высокие затраты</h4>
                <div class="cell-content">
                    <p>Стратегические проекты</p>
                </div>
            </div>
            <div class="matrix-cell high-impact-low-effort">
                <h4>Высокое влияние, Низкие затраты</h4>
                <div class="cell-content">
                    <p>Быстрые победы</p>
                </div>
            </div>
            <div class="matrix-cell low-impact-high-effort">
                <h4>Низкое влияние, Высокие затраты</h4>
                <div class="cell-content">
                    <p>Избегать</p>
                </div>
            </div>
            <div class="matrix-cell low-impact-low-effort">
                <h4>Низкое влияние, Низкие затраты</h4>
                <div class="cell-content">
                    <p>Мелкие улучшения</p>
                </div>
            </div>
        </div>'''
        
        matrix_html += '</div>\n</section>\n'
        return matrix_html
    
    def _generate_roadmap_section(self, defense: Dict[str, Any]) -> str:
        """Секция дорожной карты"""
        roadmap_html = '<section class="roadmap-section">\n'
        roadmap_html += '<h2>🗺️ Дорожная карта внедрения</h2>\n'
        roadmap_html += '<div class="roadmap-timeline">\n'
        
        # Создание временной шкалы
        phases = [
            {'name': 'Фаза 1', 'duration': '1-2 месяца', 'focus': 'Быстрые победы', 'description': 'Внедрение простых улучшений с высоким влиянием'},
            {'name': 'Фаза 2', 'duration': '3-6 месяцев', 'focus': 'Средние проекты', 'description': 'Реализация средних по сложности улучшений'},
            {'name': 'Фаза 3', 'duration': '6-12 месяцев', 'focus': 'Стратегические проекты', 'description': 'Крупные изменения продукта'}
        ]
        
        for phase in phases:
            roadmap_html += f'''
            <div class="roadmap-phase">
                <div class="phase-header">
                    <h3>{phase['name']}</h3>
                    <span class="phase-duration">{phase['duration']}</span>
                </div>
                <h4>{phase['focus']}</h4>
                <p>{phase['description']}</p>
            </div>'''
        
        roadmap_html += '</div>\n</section>\n'
        return roadmap_html
    
    def _generate_interviews_section(self, summaries: List[Any]) -> str:
        """Генерация секции с деталями по интервью"""
        if not summaries:
            return ""
        
        interviews_html = '<section class="interviews-section">\n'
        interviews_html += '<h2>📝 Детали по интервью</h2>\n'
        interviews_html += '<div class="interviews-grid">\n'
        
        for summary in summaries:
            interview_id = summary.interview_id
            pain_count = len(summary.pain_points or [])
            needs_count = len(summary.needs or [])
            quotes_count = len(summary.quotes or [])
            sentiment = summary.sentiment_score
            
            sentiment_class = 'positive' if sentiment > 0.3 else 'negative' if sentiment < -0.3 else 'neutral'
            
            interviews_html += f'''
            <div class="interview-card">
                <div class="interview-header">
                    <h3>Интервью #{interview_id}</h3>
                    <span class="sentiment-badge {sentiment_class}">{sentiment:.2f}</span>
                </div>
                <div class="interview-stats">
                    <div class="stat-item">
                        <span class="stat-value">{pain_count}</span>
                        <span class="stat-label">Боли</span>
                    </div>
                    <div class="stat-item">
                        <span class="stat-value">{needs_count}</span>
                        <span class="stat-label">Потребности</span>
                    </div>
                    <div class="stat-item">
                        <span class="stat-value">{quotes_count}</span>
                        <span class="stat-label">Цитаты</span>
                    </div>
                </div>
            </div>'''
        
        interviews_html += '</div>\n</section>\n'
        return interviews_html
    
    def _get_css_styles(self) -> str:
        """Получение CSS стилей"""
        return """
        * {
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }
        
        body {
            font-family: 'Inter', -apple-system, BlinkMacSystemFont, sans-serif;
            line-height: 1.6;
            color: #111827;
            background-color: #f9fafb;
        }
        
        .container {
            max-width: 1200px;
            margin: 0 auto;
            padding: 20px;
        }
        
        .header {
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            color: white;
            padding: 40px 20px;
            border-radius: 16px;
            margin-bottom: 30px;
            text-align: center;
        }
        
        .header h1 {
            font-size: 2.5rem;
            font-weight: 800;
            margin-bottom: 10px;
        }
        
        .subtitle {
            font-size: 1.2rem;
            opacity: 0.9;
            margin-bottom: 20px;
        }
        
        .company-info {
            margin-top: 20px;
        }
        
        .company-name {
            font-size: 1.1rem;
            font-weight: 600;
            margin-bottom: 5px;
        }
        
        .report-date {
            font-size: 0.9rem;
            opacity: 0.8;
        }
        
        section {
            background: white;
            border-radius: 12px;
            padding: 30px;
            margin-bottom: 30px;
            box-shadow: 0 1px 3px rgba(0, 0, 0, 0.1);
        }
        
        section h2 {
            font-size: 1.8rem;
            font-weight: 700;
            margin-bottom: 25px;
            color: #1f2937;
        }
        
        .metrics-grid {
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(200px, 1fr));
            gap: 20px;
        }
        
        .metric-card {
            background: linear-gradient(135deg, #f8fafc 0%, #e2e8f0 100%);
            padding: 25px;
            border-radius: 12px;
            text-align: center;
            border: 1px solid #e2e8f0;
        }
        
        .metric-value {
            font-size: 2.5rem;
            font-weight: 800;
            color: #2563eb;
            margin-bottom: 8px;
        }
        
        .metric-label {
            color: #64748b;
            font-weight: 500;
        }
        
        .insights-grid, .pain-points-grid, .recommendations-grid, .interviews-grid {
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(300px, 1fr));
            gap: 20px;
        }
        
        .insight-card, .pain-point-card, .recommendation-card, .interview-card {
            background: white;
            border: 1px solid #e5e7eb;
            border-radius: 12px;
            padding: 20px;
            transition: all 0.2s ease;
        }
        
        .insight-card:hover, .pain-point-card:hover, .recommendation-card:hover, .interview-card:hover {
            transform: translateY(-2px);
            box-shadow: 0 4px 12px rgba(0, 0, 0, 0.1);
        }
        
        .insight-header, .rec-header, .interview-header {
            display: flex;
            justify-content: space-between;
            align-items: flex-start;
            margin-bottom: 15px;
        }
        
        .insight-header h3, .rec-header h3, .interview-header h3 {
            font-size: 1.1rem;
            font-weight: 600;
            color: #1f2937;
            flex: 1;
            margin-right: 15px;
        }
        
        .severity-badge, .priority-badge, .effort-badge, .sentiment-badge {
            padding: 4px 12px;
            border-radius: 20px;
            font-size: 0.8rem;
            font-weight: 600;
            text-transform: uppercase;
        }
        
        .severity-critical { background: #fef2f2; color: #dc2626; }
        .severity-high { background: #fff7ed; color: #ea580c; }
        .severity-medium { background: #fefce8; color: #ca8a04; }
        .severity-low { background: #f0fdf4; color: #16a34a; }
        
        .priority-p0 { background: #fef2f2; color: #dc2626; }
        .priority-p1 { background: #fff7ed; color: #ea580c; }
        .priority-p2 { background: #fefce8; color: #ca8a04; }
        
        .effort-s { background: #f0fdf4; color: #16a34a; }
        .effort-m { background: #fefce8; color: #ca8a04; }
        .effort-l { background: #fff7ed; color: #ea580c; }
        .effort-xl { background: #fef2f2; color: #dc2626; }
        
        .sentiment-positive { background: #f0fdf4; color: #16a34a; }
        .sentiment-neutral { background: #fefce8; color: #ca8a04; }
        .sentiment-negative { background: #fef2f2; color: #dc2626; }
        
        .insight-description, .pain-context, .rec-description {
            color: #4b5563;
            line-height: 1.6;
            margin-bottom: 15px;
        }
        
        .insight-meta {
            display: flex;
            gap: 15px;
            font-size: 0.9rem;
        }
        
        .affected-users {
            color: #6b7280;
            font-weight: 500;
        }
        
        .rec-badges {
            display: flex;
            gap: 8px;
        }
        
        .interview-stats {
            display: grid;
            grid-template-columns: repeat(3, 1fr);
            gap: 15px;
        }
        
        .stat-item {
            text-align: center;
        }
        
        .stat-value {
            display: block;
            font-size: 1.5rem;
            font-weight: 700;
            color: #2563eb;
        }
        
        .stat-label {
            font-size: 0.8rem;
            color: #6b7280;
            text-transform: uppercase;
            font-weight: 500;
        }
        
        .severity-group {
            margin-bottom: 30px;
        }
        
        .severity-title {
            font-size: 1.3rem;
            font-weight: 600;
            margin-bottom: 15px;
            padding: 10px 0;
            border-bottom: 2px solid;
        }
        
        .severity-title.critical { color: #dc2626; border-color: #dc2626; }
        .severity-title.high { color: #ea580c; border-color: #ea580c; }
        .severity-title.medium { color: #ca8a04; border-color: #ca8a04; }
        .severity-title.low { color: #16a34a; border-color: #16a34a; }
        
        .footer {
            text-align: center;
            padding: 30px;
            color: #6b7280;
            font-size: 0.9rem;
        }
        
        .footer p {
            margin-bottom: 8px;
        }
        
        @media (max-width: 768px) {
            .container {
                padding: 15px;
            }
            
            .header {
                padding: 30px 15px;
            }
            
            .header h1 {
                font-size: 2rem;
            }
            
            .metrics-grid {
                grid-template-columns: repeat(2, 1fr);
            }
            
            .insights-grid, .pain-points-grid, .recommendations-grid, .interviews-grid {
                grid-template-columns: 1fr;
            }
        }
        """
